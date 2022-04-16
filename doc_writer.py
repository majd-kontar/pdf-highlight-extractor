from docx import Document


class DocWriter:
    def __init__(self):
        self.document = Document()

    def write(self, filename: str, highlighted_text: list):
        """
        Writes the highlighted text to a doc-file under a heading of the pdf name
        :param filename: pdf name
        :param highlighted_text: highlighted text in pdf
        :return: None
        """
        self.document.add_heading(filename, 1)
        for text in highlighted_text:
            self.document.add_paragraph(text, style='List Bullet')

    def save(self, filepath):
        """
        Saves the doc at the given path
        :param filepath: path where the doc is to be saved along with the name and docx extension
        :return: None
        """
        self.document.save(filepath)


if __name__ == '__main__':
    text = ['Lorem ipsum dolor sit amet, consectetuer adipiscing elit.',
            'Aliquam tincidunt mauris eu risus.',
            'Vestibulum auctor dapibus neque.',
            'Nunc dignissim risus id metus.',
            'Cras ornare tristique elit.',
            'Vivamus vestibulum ntulla nec ante.',
            'Praesent placerat risus quis eros.',
            'Fusce pellentesque suscipit nibh.',
            'Integer vitae libero ac risus egestas placerat.',
            'Vestibulum commodo felis quis tortor.',
            'Ut aliquam sollicitudin leo.',
            'Cras iaculis ultricies nulla.',
            'Donec quis dui at dolor tempor interdum.']
    doc_writer = DocWriter()
    doc_writer.write('Filename', text)
    doc_writer.save("test.docx")
